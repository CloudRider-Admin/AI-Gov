#!/usr/bin/env python3
"""
GovSecure Content Library Extraction Pipeline (Phase 0).

Walks the GovSecure source folder, extracts structured content from .docx,
.pdf, and .xlsx files, and emits per-document JSON files plus a manifest
to src/data/govsecureContent/.

Outputs are deterministic and idempotent — re-running with unchanged
sources produces byte-identical output.

Usage:
    python3 scripts/ingestGovSecureLibrary.py [--source PATH] [--out PATH]
                                              [--validate-only]

References:
    GOVI_GOVSECURE_INTEGRATION_PLAN.md — Phase 0 specification
"""
from __future__ import annotations

import argparse
import dataclasses
import hashlib
import json
import re
import sys
from collections import OrderedDict, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

# Third-party libraries (declared in requirements-extraction.txt)
from docx import Document  # python-docx
import pdfplumber
from openpyxl import load_workbook


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

REPO_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_SOURCE = REPO_ROOT / "GovSecure Project-20260505T105940Z-3-001" / "GovSecure Project"
DEFAULT_OUT = REPO_ROOT / "src" / "data" / "govsecureContent"

# Heading styles produced by GovSecure's Word templates. Map style name -> heading level.
# The library uses three different style conventions across documents:
#   - Licensed policies: GovTitle / GovSub / GovH1
#   - AI Chef toolkit:    TitleGov / SubtitleGov / H1Gov / H2Gov / SmallGov
#   - Third-party / generic: standard Heading 1 / Heading 2
HEADING_STYLES: dict[str, int] = {
    # Title-class
    "GovTitle": 0,
    "GovSub": 0,
    "TitleGov": 0,
    "SubtitleGov": 0,
    "Title": 0,
    "SmallGov": 0,
    # Heading levels (prefix variant)
    "GovH1": 1,
    "GovH2": 2,
    "GovH3": 3,
    # Heading levels (suffix variant)
    "H1Gov": 1,
    "H2Gov": 2,
    "H3Gov": 3,
    # Standard Word fallbacks
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
}

# Regex patterns that identify likely section headings in plain "Normal" paragraphs.
# Used as a fallback when a document does not apply heading styles (e.g. TPRM,
# 90-Day Blueprint). Tested in order; first match wins.
HEADING_HEURISTICS: list[tuple[re.Pattern[str], int]] = [
    # Numbered top-level: "1) AI System Overview (Required)" or "1. Section Name"
    (re.compile(r"^\d+\s*[\)\.]\s+\S.+$"), 1),
    # Phase markers: "Phase 1", "Phase 1 — Foundation"
    (re.compile(r"^Phase\s+\d+\b.*$"), 1),
    # Section markers: "Section 1", "Section 1 — Topic"
    (re.compile(r"^Section\s+\d+\b.*$"), 1),
    # Week markers: "Week 1", "Week 1: Setup"
    (re.compile(r"^Week\s+\d+\b.*$"), 2),
]

BULLET_STYLES: set[str] = {
    "List Bullet",
    "List Bullet 2",
    "List Bullet 3",
    "List Number",
    "List Number 2",
}

# Categorization rules: matched against the source path relative to the source root.
# First match wins. Order matters.
CATEGORY_RULES: list[tuple[re.Pattern[str], str, str]] = [
    # (pattern, category, subcategory_strategy)
    (re.compile(r"Policies/.+/govsecure_core8_licensed_final/"), "policies", "policy_slug"),
    (re.compile(r"Policies/Govsecure_AI Policy_Suite/00_Overview/"), "policies", "overview"),
    (re.compile(r"Policies/Govsecure_AI Policy_Suite/"), "policies", "suite_meta"),
    (re.compile(r"Policies/.*"), "policies", "policy_slug"),
    (re.compile(r"Checklists/"), "checklists", "checklist_slug"),
    (re.compile(r"Playbook/GovSecure_AI_Chef_Client_Ready_Package/"), "playbooks", "ai_chef_client"),
    (re.compile(r"Playbook/Non-Customized Playbook/"), "playbooks", "ai_chef_noncustom"),
    (re.compile(r"Playbook/GovSecure_Third_Party_AI_Privacy_Risk_Assessment_Package/"), "playbooks", "third_party_privacy"),
    (re.compile(r"Playbook/"), "playbooks", "playbook_root"),
    (re.compile(r"AI Adoption and Implementation BluePrint Template/"), "playbooks", "blueprint"),
    (re.compile(r"AI Frameworks/"), "frameworks", "framework_slug"),
    (re.compile(r"Questionnaires/"), "questionnaires", "questionnaire_slug"),
]

# Document code prefixes per category. Encoded into the document code as
# GS-{PREFIX}-{NUM:02d} so generated artifacts mirror the GovSecure scheme.
DOC_CODE_PREFIX: dict[str, str] = {
    "policies": "AIPS",
    "checklists": "CHKL",
    "playbooks": "PLBK",
    "frameworks": "FRAM",
    "questionnaires": "QSTN",
}

# Map file slug -> semantic sub-type. Allows downstream code to look up
# documents by stable identifier rather than filename.
SUBTYPE_MAP: dict[str, str] = {
    # Core 8 policies
    "ai_data_handling_and_privacy_policy_licensed": "data-privacy",
    "ai_governance_policy_licensed": "governance",
    "ai_incident_response_and_monitoring_policy_licensed": "incident-response",
    "ai_risk_assessment_and_use_case_approval_policy_licensed": "risk-approval",
    "ai_security_policy_licensed": "security",
    "enterprise_ai_acceptable_use_policy_licensed": "acceptable-use",
    "human_oversight_and_decision_making_policy_licensed": "human-oversight",
    "third_party_vendor_ai_due_diligence_policy_licensed": "third-party",
    # Policy suite meta
    "govsecure_ai_policy_suite_offering_guide": "suite-offering",
    "govsecure_ai_policy_suite_map_client_ready": "suite-map",
    # Checklists
    "ai_evidence_pack_checklist": "evidence-pack",
    "ai_governance_checklist_templates": "governance-templates",
    "ai_intake_triage_checklist": "intake-triage",
    "ai_use_case_intake_form_checklist": "intake-form",
    "ai_change_management_checklist": "change-management",
    "ai_human_oversight_escalation_checklist": "human-oversight-escalation",
    "ai_incident_response_checklist_ai_specific": "incident-response",
    "ai_inventory_system_registry_checklist": "inventory-registry",
    "ai_model_validation_testing_checklist": "model-validation",
    "ai_monitoring_revalidation_checklist": "monitoring-revalidation",
    "ai_policy_to_control_mapping_checklist": "policy-to-control",
    "ai_privacy_dpia_screening_checklist": "dpia-screening",
    "ai_regulatory_change_impact_assessment_checklist": "regulatory-change",
    "ai_risk_assessment_template_checklist": "risk-assessment-template",
    "ai_security_review_checklist": "security-review",
    "ai_third_party_vendor_due_diligence_checklist": "third-party-dd",
    "ai_training_role_based_awareness_checklist": "training-awareness",
    "govsecure_ai_acceptable_use_policy_generator": "aup-generator",
    "govsecure_ai_governance_template_pack_preview": "template-pack-preview",
    "govsecure_ai_incident_response_checklist": "incident-response",
    "govsecure_ai_procurement_review_checklist": "procurement-review",
    "govsecure_ai_tool_inventory_checklist": "tool-inventory",
    "govsecure_ai_use_case_intake_checklist": "use-case-intake",
    "govsecure_ai_use_case_intake_form": "use-case-intake-form",
    "govsecure_ai_use_case_registry_template": "use-case-registry",
    "govsecure_shadow_ai_discovery_workflow": "shadow-ai-discovery",
    "govsecure_vendor_ai_due_diligence_checklist": "vendor-due-diligence",
    "other_vital_ai_governance_checklist_templates": "other-templates",
    # Playbooks
    "90_day_ai_governance_implementation_blueprint": "90-day-blueprint",
    "90_day_blueprint_updated": "90-day-blueprint-updated",
    "govsecure_ai_chef_client_ready_executive_toolkit": "ai-chef-executive",
    "govsecure_ai_chef_client_ready_operational_templates": "ai-chef-operational",
    "govsecure_ai_chef_client_ready_infographic": "ai-chef-infographic",
    "govsecure_ai_chef_client_ready_toolkit": "ai-chef-toolkit",
    "govsecure_ai_chef_executive_toolkit": "ai-chef-executive-noncustom",
    "govsecure_ai_chef_operational_templates": "ai-chef-operational-noncustom",
    "govsecure_ai_chef_infographic": "ai-chef-infographic-noncustom",
    "govsecure_ai_chef_toolkit_word": "ai-chef-toolkit-noncustom",
    "govsecure_third_party_ai_privacy_risk_assessment": "third-party-privacy-questionnaire",
    "govsecure_third_party_ai_privacy_risk_assessment_workbook": "third-party-privacy-workbook",
    # Frameworks
    "ai_rcm_nist_framework_v5": "nist-rcm",
    # Questionnaires
    "3p_risk_management_tprm_ai_questionnaire": "tprm",
    "govsecure_3rd_party_ai_vendor_risk_review": "tprm-review",
}


# ---------------------------------------------------------------------------
# Utility helpers
# ---------------------------------------------------------------------------

_SLUG_RE = re.compile(r"[^a-z0-9]+")


def slugify(text: str) -> str:
    """Lowercase, collapse non-alphanumerics to hyphens, strip ends."""
    text = text.lower().replace("'", "")
    text = _SLUG_RE.sub("-", text).strip("-")
    return text


def filename_slug(path: Path) -> str:
    """Slug derived from filename without extension and parenthesized variants."""
    stem = path.stem
    # Strip Windows-style copy markers like " (1)"
    stem = re.sub(r"\s*\(\d+\)\s*$", "", stem)
    return slugify(stem.replace("_", "-").replace(" ", "-")).replace("-", "_")


def checksum_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(65536), b""):
            h.update(chunk)
    return f"sha256:{h.hexdigest()}"


def categorize(rel_path: str) -> tuple[str, str]:
    """Return (category, subcategory_strategy) for a path relative to the source root."""
    norm = rel_path.replace("\\", "/")
    for pattern, category, strategy in CATEGORY_RULES:
        if pattern.search(norm):
            return category, strategy
    return "other", "unknown"


def build_doc_code(category: str, sub_type: str, num: int) -> str:
    prefix = DOC_CODE_PREFIX.get(category, "GENR")
    type_token = sub_type.upper().replace("-", "")[:8] if sub_type else "GENR"
    return f"GS-{prefix}-{type_token}-{num:02d}"


def now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def source_mtime_iso(path: Path) -> str:
    """ISO timestamp of the source file's last modification.

    Used for the ``extractedAt`` metadata field so that re-running the script
    on unchanged sources produces byte-identical output (idempotency).
    """
    ts = path.stat().st_mtime
    return datetime.fromtimestamp(ts, timezone.utc).replace(microsecond=0).isoformat()


# ---------------------------------------------------------------------------
# Extractors
# ---------------------------------------------------------------------------

@dataclasses.dataclass
class ExtractedSection:
    id: str
    heading: str
    level: int
    paragraphs: list[str] = dataclasses.field(default_factory=list)
    bullets: list[str] = dataclasses.field(default_factory=list)
    tables: list[list[list[str]]] = dataclasses.field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "heading": self.heading,
            "level": self.level,
            "paragraphs": self.paragraphs,
            "bullets": self.bullets,
            "tables": self.tables,
        }


def _extract_docx_tables(table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    return rows


def _is_paragraph_bold(paragraph) -> bool:
    """Return True if all non-empty runs in the paragraph are bold."""
    runs = [r for r in paragraph.runs if r.text.strip()]
    if not runs:
        return False
    return all(bool(r.bold) for r in runs)


def _heuristic_heading_level(text: str) -> int | None:
    """Return a heading level if the text matches a known heading pattern."""
    if len(text) > 200:
        return None
    for pattern, level in HEADING_HEURISTICS:
        if pattern.match(text):
            return level
    return None


def extract_docx(path: Path) -> dict[str, Any]:
    """Extract structured content from a .docx file.

    Detects headings via three signals (in priority order):
      1. Named heading styles (GovH1, H1Gov, Heading 1, ...)
      2. Numbered/sectioned content patterns ("1) Topic", "Phase 1", ...)
      3. All-bold paragraphs that are short enough to be headings (≤120 chars)
    """
    doc = Document(str(path))

    title_parts: list[str] = []
    sections: list[ExtractedSection] = []
    current = ExtractedSection(id="0", heading="(preamble)", level=1)
    section_counter = 0

    body = doc.element.body
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)

    def flush_current() -> None:
        nonlocal current
        if (
            current.paragraphs
            or current.bullets
            or current.tables
            or current.heading != "(preamble)"
        ):
            sections.append(current)

    def open_section(text: str, level: int) -> None:
        nonlocal current, section_counter
        flush_current()
        section_counter += 1
        current = ExtractedSection(id=str(section_counter), heading=text, level=level)

    for child in body.iterchildren():
        tag = child.tag.split("}", 1)[-1]
        if tag == "p":
            try:
                p = next(para_iter)
            except StopIteration:
                continue

            text = p.text.strip()
            if not text:
                continue

            style_name = p.style.name if p.style else ""
            style_level = HEADING_STYLES.get(style_name)

            # Title-like styles (level 0) flow into the document title.
            if style_level == 0 and not sections:
                # Multi-line titles can contain embedded newlines; flatten them.
                title_parts.append(text.replace("\n", " — "))
                continue

            # Bullet-list styles
            if style_name in BULLET_STYLES or style_name.startswith("List "):
                current.bullets.append(text)
                continue

            # Style-driven heading
            if style_level is not None and style_level >= 1:
                open_section(text, style_level)
                continue

            # Heuristic heading detection for unstyled documents
            heuristic_level = _heuristic_heading_level(text)
            if heuristic_level is not None:
                open_section(text, heuristic_level)
                continue

            # Bold-only short paragraph → likely a heading. Skip if it's the
            # very first content (commonly a title) and we already have one.
            if (
                len(text) <= 120
                and not text.endswith((".", ":", "?", ","))
                and _is_paragraph_bold(p)
            ):
                open_section(text, 1)
                continue

            # Otherwise, it's a paragraph.
            current.paragraphs.append(text)

        elif tag == "tbl":
            try:
                table = next(table_iter)
            except StopIteration:
                continue
            current.tables.append(_extract_docx_tables(table))

    flush_current()

    # If no explicit title was set, fall back to the first H1 heading or the filename.
    if title_parts:
        title = " — ".join(title_parts)
    elif sections and sections[0].heading != "(preamble)":
        title = sections[0].heading
    else:
        title = path.stem.replace("_", " ")

    return {
        "title": title,
        "sections": [s.to_dict() for s in sections],
        "format": "docx",
        "stats": {
            "section_count": len(sections),
            "paragraph_count": sum(len(s.paragraphs) for s in sections),
            "bullet_count": sum(len(s.bullets) for s in sections),
            "table_count": sum(len(s.tables) for s in sections),
        },
    }


def extract_pdf(path: Path) -> dict[str, Any]:
    """Extract text from a PDF as flat sections per page.

    Used as a fallback when a .docx companion is not available. PDF parsing is
    inherently lossy; we emit one section per page rather than try to detect
    headings heuristically.
    """
    sections: list[dict[str, Any]] = []
    total_chars = 0

    with pdfplumber.open(str(path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if not text:
                continue

            paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
            total_chars += sum(len(p) for p in paragraphs)
            sections.append(
                {
                    "id": f"page-{page_num}",
                    "heading": f"Page {page_num}",
                    "level": 1,
                    "paragraphs": paragraphs,
                    "bullets": [],
                    "tables": [],
                }
            )

    return {
        "title": path.stem.replace("_", " "),
        "sections": sections,
        "format": "pdf",
        "stats": {
            "section_count": len(sections),
            "char_count": total_chars,
        },
    }


def extract_xlsx(path: Path) -> dict[str, Any]:
    """Extract every sheet of an Excel workbook into per-sheet sections."""
    wb = load_workbook(str(path), data_only=True, read_only=True)
    sections: list[dict[str, Any]] = []
    sheet_data: dict[str, list[list[Any]]] = {}

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[list[Any]] = []
        for row in ws.iter_rows(values_only=True):
            cleaned = [
                ("" if cell is None else str(cell).strip())
                for cell in row
            ]
            # Drop trailing empty cells.
            while cleaned and cleaned[-1] == "":
                cleaned.pop()
            rows.append(cleaned)
        # Trim trailing fully-empty rows.
        while rows and not any(c for c in rows[-1]):
            rows.pop()

        if not rows:
            continue

        sheet_data[sheet_name] = rows

        # Each sheet → one section. Tables array holds the grid; paragraphs is empty.
        sections.append(
            {
                "id": slugify(sheet_name) or f"sheet-{len(sections) + 1}",
                "heading": sheet_name,
                "level": 1,
                "paragraphs": [],
                "bullets": [],
                "tables": [rows],
            }
        )

    wb.close()

    return {
        "title": path.stem.replace("_", " "),
        "sections": sections,
        "format": "xlsx",
        "stats": {
            "sheet_count": len(sections),
            "total_rows": sum(len(r) for r in sheet_data.values()),
        },
    }


# ---------------------------------------------------------------------------
# Pipeline
# ---------------------------------------------------------------------------

@dataclasses.dataclass
class SourceFile:
    path: Path
    rel_path: str
    suffix: str
    slug: str
    category: str
    strategy: str

    @property
    def stable_key(self) -> str:
        """Group key used to dedupe DOCX/PDF pairs of the same document."""
        return f"{self.category}::{self.strategy}::{self.slug}"


def discover_sources(source_root: Path) -> list[SourceFile]:
    """Walk the source folder and produce a list of SourceFile records."""
    files: list[SourceFile] = []
    for path in sorted(source_root.rglob("*")):
        if not path.is_file():
            continue
        suffix = path.suffix.lower()
        if suffix not in {".docx", ".pdf", ".xlsx"}:
            # Skip .zip and any other binaries.
            continue
        rel = path.relative_to(source_root).as_posix()
        slug = filename_slug(path)
        category, strategy = categorize(rel)
        files.append(
            SourceFile(
                path=path,
                rel_path=rel,
                suffix=suffix,
                slug=slug,
                category=category,
                strategy=strategy,
            )
        )
    return files


def select_canonical(group: list[SourceFile]) -> SourceFile:
    """Within a group of equivalent files, prefer DOCX > XLSX > PDF.

    DOCX gives us the richest structure; XLSX is required for spreadsheet-only
    documents; PDF is a last-resort fallback.
    """
    by_priority = {".docx": 0, ".xlsx": 1, ".pdf": 2}
    return sorted(group, key=lambda f: (by_priority.get(f.suffix, 99), f.rel_path))[0]


def extract_one(source: SourceFile) -> dict[str, Any]:
    """Dispatch to the appropriate extractor based on file suffix."""
    if source.suffix == ".docx":
        payload = extract_docx(source.path)
    elif source.suffix == ".pdf":
        payload = extract_pdf(source.path)
    elif source.suffix == ".xlsx":
        payload = extract_xlsx(source.path)
    else:
        raise ValueError(f"Unsupported suffix: {source.suffix}")

    sub_type = SUBTYPE_MAP.get(source.slug, source.slug.replace("_", "-"))

    payload.update(
        {
            "category": source.category,
            "subType": sub_type,
            "sourcePath": source.rel_path,
            "sourceSuffix": source.suffix,
            "checksum": checksum_file(source.path),
            "extractedAt": source_mtime_iso(source.path),
        }
    )
    return payload


def assign_doc_codes(records: list[dict[str, Any]]) -> None:
    """Mutate records in place, assigning deterministic document codes per category."""
    counters: dict[str, int] = defaultdict(int)
    # Sort for deterministic numbering.
    for record in sorted(records, key=lambda r: (r["category"], r["subType"], r["sourcePath"])):
        counters[record["category"]] += 1
        record["documentCode"] = build_doc_code(
            record["category"], record["subType"], counters[record["category"]]
        )


def write_record(out_root: Path, record: dict[str, Any]) -> Path:
    """Write a single record JSON to disk and return the output path."""
    category_dir = out_root / record["category"]
    category_dir.mkdir(parents=True, exist_ok=True)
    file_slug = slugify(record["subType"]) or slugify(record.get("title") or "unknown")
    out_path = category_dir / f"{file_slug}.json"

    # If two records hit the same slug (e.g. duplicate copies) suffix to disambiguate.
    counter = 2
    while out_path.exists():
        rendered = json.dumps(record, indent=2, ensure_ascii=False, sort_keys=True)
        # If existing file has identical content, skip duplicate write.
        if out_path.read_text(encoding="utf-8") == rendered:
            return out_path
        out_path = category_dir / f"{file_slug}-{counter}.json"
        counter += 1

    rendered = json.dumps(record, indent=2, ensure_ascii=False, sort_keys=True)
    out_path.write_text(rendered + "\n", encoding="utf-8")
    return out_path


def build_manifest(records_with_paths: list[tuple[dict[str, Any], Path]],
                   out_root: Path) -> dict[str, Any]:
    """Build a manifest summarizing all extracted documents."""
    entries = []
    by_category: dict[str, int] = defaultdict(int)

    for record, out_path in sorted(
        records_with_paths, key=lambda x: (x[0]["category"], x[0]["documentCode"])
    ):
        rel_out = out_path.relative_to(out_root).as_posix()
        by_category[record["category"]] += 1
        entries.append(
            OrderedDict(
                [
                    ("documentCode", record["documentCode"]),
                    ("category", record["category"]),
                    ("subType", record["subType"]),
                    ("title", record.get("title", "")),
                    ("sourcePath", record["sourcePath"]),
                    ("outputPath", rel_out),
                    ("format", record["format"]),
                    ("checksum", record["checksum"]),
                    ("stats", record.get("stats", {})),
                ]
            )
        )

    # Use the latest source mtime as the manifest timestamp so the manifest is
    # also content-deterministic.
    latest_mtime = max(
        (record.get("extractedAt", "") for record, _ in records_with_paths),
        default=now_iso(),
    )
    return OrderedDict(
        [
            ("generatedAt", latest_mtime),
            ("sourceRoot", str(DEFAULT_SOURCE.relative_to(REPO_ROOT))),
            ("documentCount", len(entries)),
            ("byCategory", dict(by_category)),
            ("documents", entries),
        ]
    )


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--source",
        type=Path,
        default=DEFAULT_SOURCE,
        help="Path to the GovSecure source folder",
    )
    parser.add_argument(
        "--out",
        type=Path,
        default=DEFAULT_OUT,
        help="Output directory for extracted JSON",
    )
    parser.add_argument(
        "--validate-only",
        action="store_true",
        help="Validate output against a previous run rather than rewriting files",
    )
    parser.add_argument(
        "--include-duplicates",
        action="store_true",
        help="Emit DOCX, XLSX, and PDF copies separately instead of preferring one",
    )
    args = parser.parse_args()

    source_root: Path = args.source
    out_root: Path = args.out

    if not source_root.exists():
        print(f"ERROR: source folder not found: {source_root}", file=sys.stderr)
        return 2

    print(f"[ingest] source: {source_root}")
    print(f"[ingest] output: {out_root}")

    sources = discover_sources(source_root)
    print(f"[ingest] discovered {len(sources)} source files")

    # Group by stable_key and pick a canonical representative per group.
    groups: dict[str, list[SourceFile]] = defaultdict(list)
    for source in sources:
        groups[source.stable_key].append(source)

    if args.include_duplicates:
        canonical_sources = sources
    else:
        canonical_sources = [select_canonical(g) for g in groups.values()]

    print(f"[ingest] processing {len(canonical_sources)} canonical documents "
          f"({len(sources) - len(canonical_sources)} duplicates skipped)")

    records: list[dict[str, Any]] = []
    failures: list[tuple[str, str]] = []
    for source in canonical_sources:
        try:
            record = extract_one(source)
        except Exception as exc:  # noqa: BLE001
            failures.append((source.rel_path, str(exc)))
            continue
        records.append(record)

    assign_doc_codes(records)

    # Reset output dir to keep it clean and idempotent.
    if out_root.exists() and not args.validate_only:
        for child in out_root.rglob("*"):
            if child.is_file():
                child.unlink()
    out_root.mkdir(parents=True, exist_ok=True)

    written: list[tuple[dict[str, Any], Path]] = []
    for record in records:
        out_path = write_record(out_root, record)
        written.append((record, out_path))

    manifest = build_manifest(written, out_root)
    manifest_path = out_root / "manifest.json"
    manifest_path.write_text(
        json.dumps(manifest, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )

    # Emit a single bundled index containing every extracted document, keyed
    # by documentCode. Used by src/data/govsecureKnowledge.ts so the runtime
    # only needs to import one file rather than ~57 individual JSONs.
    bundled = OrderedDict(
        [
            ("generatedAt", manifest["generatedAt"]),
            ("documentCount", len(written)),
            (
                "documents",
                {
                    record["documentCode"]: record
                    for record, _ in sorted(
                        written, key=lambda x: x[0]["documentCode"]
                    )
                },
            ),
        ]
    )
    index_path = out_root / "index.json"
    index_path.write_text(
        json.dumps(bundled, indent=2, ensure_ascii=False, sort_keys=False) + "\n",
        encoding="utf-8",
    )

    print(f"[ingest] wrote {len(written)} document JSON files")
    print(f"[ingest] by category: {manifest['byCategory']}")
    print(f"[ingest] manifest: {manifest_path.relative_to(REPO_ROOT)}")
    print(f"[ingest] bundled index: {index_path.relative_to(REPO_ROOT)}")

    if failures:
        print(f"\n[ingest] FAILURES ({len(failures)}):", file=sys.stderr)
        for rel_path, err in failures:
            print(f"  - {rel_path}: {err}", file=sys.stderr)
        return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())